#!/usr/bin/env python3
"""
Parse 分子生物学历年卷.docx → public/question-bank.json

Question types:
  - translation:    中英名词互译 (~104, from 7 tables)
  - true-false:     判断题 (~14)
  - multiple-choice: 选择题 (~18)
  - short-answer:   简答题 (~45)
  - essay:          分析论述题 (~4)
"""

import json
import re
import sys
from datetime import datetime
from pathlib import Path
from typing import Any

try:
    from docx import Document
except ImportError:
    print("Error: python-docx required. Run: pip install python-docx")
    sys.exit(1)

ROOT = Path(__file__).resolve().parent.parent
DOCX_PATH = ROOT / "data" / "分子生物学历年卷.docx"
OUTPUT_PATH = ROOT / "public" / "question-bank.json"


# ── helpers ──────────────────────────────────────────────────────────

def clean(text: str) -> str:
    return re.sub(r'\s+', ' ', text).strip()


def first_line(text: str) -> str:
    """Return the first non-empty line of a multi-line paragraph."""
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    return lines[0] if lines else ''


def remaining_lines(text: str) -> str:
    """Return everything after the first line."""
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    return '\n'.join(lines[1:]) if len(lines) > 1 else ''


def parse_abbreviation(en_term: str) -> tuple[str, str | None]:
    """Parse English term for abbreviation pattern.
    "ChIP (Chromatin immunoprecipitation)" → ("ChIP", "Chromatin immunoprecipitation")
    """
    m = re.match(r'^(.+?)\s*[（(]([^）)]+?)[）)]$', en_term)
    if m:
        a = clean(m.group(1))
        b = clean(m.group(2))
        return (a, b) if len(a) <= len(b) else (b, a)
    return en_term.strip(), None


def build_acceptable_answers(answer_term: str, answer_full_term: str | None) -> list[str]:
    answers = [answer_term.strip().lower()]
    if answer_full_term:
        ft = answer_full_term.strip().lower()
        if ft != answers[0]:
            answers.append(ft)
        combined = f"{answer_term.strip()} ({answer_full_term.strip()})".lower()
        if combined not in answers:
            answers.append(combined)
    return answers


def is_subcategory_header(text: str) -> bool:
    return bool(re.match(r'^[（(][一二三四五六七八九十]+[）)]\s*', text))


def is_chapter_header(text: str) -> bool:
    return bool(re.match(r'^[一二三四五六七八九十]、', text))


def is_answer_line(text: str) -> bool:
    """Check if a line looks like answer content (bullet, numbered, etc.)."""
    if re.match(r'^[①②③④⑤⑥⑦⑧⑨⑩]', text):
        return True
    if re.match(r'^\d+[\.\、\)]', text):
        return True
    if text.startswith('? ') or text.startswith('? '):
        return True
    if text.startswith('- ') or text.startswith('• '):
        return True
    return False


# ── section boundary detection ──────────────────────────────────────

def find_section_boundaries(para_map: dict[int, str]) -> dict[str, tuple[int, int] | None]:
    """Find start/end paragraph indices for each question type section."""
    section_starts: dict[str, int] = {}

    for idx, text in para_map.items():
        if not is_chapter_header(text):
            continue
        if '判断题' in text:
            section_starts['true-false'] = idx
        elif '选择题' in text:
            section_starts['multiple-choice'] = idx
        elif '简答题' in text:
            section_starts['short-answer'] = idx
        elif ('分析' in text or '论述' in text) and ('论述' in text or '论' in text):
            section_starts['essay'] = idx

    sorted_sections = sorted(section_starts.items(), key=lambda x: x[1])
    max_idx = max(para_map.keys())

    boundaries: dict[str, tuple[int, int] | None] = {
        'true-false': None,
        'multiple-choice': None,
        'short-answer': None,
        'essay': None,
    }

    for i, (section_name, start_idx) in enumerate(sorted_sections):
        if i < len(sorted_sections) - 1:
            end_idx = sorted_sections[i + 1][1] - 1
        else:
            end_idx = max_idx
        boundaries[section_name] = (start_idx, end_idx)

    return boundaries


# ── parsers ──────────────────────────────────────────────────────────

def parse_translation_tables(doc: Any) -> list[dict]:
    subcategory_names = [
        "基础分子生物学概念", "核酸结构与功能", "蛋白质结构与修饰",
        "表观遗传学", "分子克隆与实验技术", "基因组与遗传学", "其他",
    ]

    questions = []
    question_num = 0

    for table_idx in range(min(7, len(doc.tables))):
        table = doc.tables[table_idx]
        subcat = subcategory_names[table_idx] if table_idx < len(subcategory_names) else f"子分类{table_idx+1}"
        category_id = f"translation-{table_idx + 1}"

        for row in table.rows[1:]:
            cells = [clean(c.text) for c in row.cells]
            if len(cells) < 2:
                continue

            # Tables 1, 4, 5 have 3 columns where col 2 is the true Chinese translation
            # (col 1 is sometimes a duplicate of the English term instead of Chinese)
            zh_term = cells[2] if len(cells) >= 3 else cells[1]
            en_term = cells[0]

            if not en_term or not zh_term:
                continue

            en_variants = [clean(e) for e in en_term.split('/') if clean(e)]
            primary_en = en_variants[0]

            answer_term, answer_full_term = parse_abbreviation(primary_en)
            acceptable_answers = build_acceptable_answers(answer_term, answer_full_term)

            for variant in en_variants[1:]:
                var_term, var_full = parse_abbreviation(variant)
                acceptable_answers.append(var_term.strip().lower())
                if var_full:
                    acceptable_answers.append(var_full.strip().lower())

            question_num += 1
            questions.append({
                "id": f"{category_id}-q-{question_num:03d}",
                "type": "translation",
                "categoryId": category_id,
                "categoryTitle": subcat,
                "parentTitle": "中英名词互译",
                "number": question_num,
                "prompt": zh_term,           # show Chinese, ask for English
                "direction": "zh-to-en",
                "answerTerm": answer_term,
                "answerFullTerm": answer_full_term,
                "acceptableAnswers": list(set(acceptable_answers)),
                "chineseMeaning": zh_term,
            })

    return questions


def parse_true_false(doc: Any, start_idx: int, end_idx: int) -> list[dict]:
    questions = []
    question_num = 0
    current_subcat = ""
    current_subcat_id = ""
    subcat_counter = 0
    paragraphs = doc.paragraphs
    i = start_idx + 1

    while i <= end_idx:
        text = clean(paragraphs[i].text)
        if not text:
            i += 1
            continue

        if is_subcategory_header(text):
            subcat_counter += 1
            current_subcat = text
            current_subcat_id = f"true-false-{subcat_counter}"
            i += 1
            continue
        if is_chapter_header(text):
            i += 1
            continue

        if text and not text.startswith('结论') and not text.startswith('解析'):
            statement = text
            conclusion = ""
            explanation = ""
            j = i + 1
            while j <= min(i + 5, end_idx):
                look_text = clean(paragraphs[j].text)
                if look_text.startswith('结论：') or look_text.startswith('结论:'):
                    conclusion = clean(look_text.replace('结论：', '').replace('结论:', ''))
                elif look_text.startswith('解析：') or look_text.startswith('解析:'):
                    explanation = clean(look_text.replace('解析：', '').replace('解析:', ''))
                elif look_text and not (look_text.startswith('结论') or look_text.startswith('解析')):
                    if conclusion:
                        break
                j += 1

            if statement and conclusion:
                question_num += 1
                questions.append({
                    "id": f"{current_subcat_id}-q-{question_num:03d}" if current_subcat_id else f"tf-q-{question_num:03d}",
                    "type": "true-false",
                    "categoryId": current_subcat_id or "true-false-1",
                    "categoryTitle": current_subcat or "判断题",
                    "parentTitle": "判断题",
                    "number": question_num,
                    "prompt": statement,
                    "answerIsTrue": '正确' in conclusion or 'TRUE' in conclusion.upper(),
                    "explanation": explanation,
                })
            i = j - 1
        i += 1

    return questions


def parse_multiple_choice(doc: Any, start_idx: int, end_idx: int) -> list[dict]:
    questions = []
    question_num = 0
    current_subcat = ""
    current_subcat_id = ""
    subcat_counter = 0
    paragraphs = doc.paragraphs
    i = start_idx + 1

    while i <= end_idx:
        text = clean(paragraphs[i].text)
        if not text:
            i += 1
            continue

        if is_subcategory_header(text):
            subcat_counter += 1
            current_subcat = text
            current_subcat_id = f"multiple-choice-{subcat_counter}"
            i += 1
            continue
        if is_chapter_header(text):
            i += 1
            continue

        if text:
            answer_text = ""
            explanation = ""
            j = i + 1
            while j <= min(i + 5, end_idx):
                look_text = clean(paragraphs[j].text)
                if look_text.startswith('答案：') or look_text.startswith('答案:'):
                    answer_text = clean(look_text.replace('答案：', '').replace('答案:', ''))
                elif look_text.startswith('解析：') or look_text.startswith('解析:'):
                    explanation = clean(look_text.replace('解析：', '').replace('解析:', ''))
                elif look_text and not (look_text.startswith('答案') or look_text.startswith('解析')):
                    if answer_text or explanation:
                        break
                j += 1

            if answer_text:
                question_num += 1
                answer_key = _extract_answer_key(answer_text)
                questions.append({
                    "id": f"{current_subcat_id}-q-{question_num:03d}" if current_subcat_id else f"mc-q-{question_num:03d}",
                    "type": "multiple-choice",
                    "categoryId": current_subcat_id or "multiple-choice-1",
                    "categoryTitle": current_subcat or "选择题",
                    "parentTitle": "选择题",
                    "number": question_num,
                    "prompt": text,
                    "options": [],
                    "answerKey": answer_key,
                    "answerText": answer_text,
                    "explanation": explanation,
                    "distractorsGenerated": False,
                })
            i = j - 1
        i += 1

    return questions


def _extract_answer_key(text: str) -> str | None:
    text = clean(text)
    if re.match(r'^[A-E]$', text):
        return text
    if re.match(r'^[A-E][、,，\s]+[A-E]', text):
        return text
    return None


def parse_qa_section(doc: Any, start_idx: int, end_idx: int,
                     question_type: str, parent_title: str,
                     type_prefix: str) -> list[dict]:
    """
    Generic parser for short-answer and essay sections.
    Uses first-line heuristic: the first line of a paragraph is the question title,
    remaining lines + following paragraphs are the answer.

    Also handles subcategories within the section.
    """
    questions = []
    question_num = 0
    current_subcat = ""
    current_subcat_id = ""
    subcat_counter = 0
    paragraphs = doc.paragraphs
    i = start_idx + 1  # skip section header

    while i <= end_idx:
        raw_text = paragraphs[i].text.strip('\n').replace('\r', '\n')
        text = clean(raw_text)

        if not text:
            i += 1
            continue

        # Check for subcategory header (first line)
        fl = first_line(raw_text)
        if is_subcategory_header(fl):
            subcat_counter += 1
            current_subcat = clean(fl)
            current_subcat_id = f"{type_prefix}-{subcat_counter}"
            i += 1
            continue

        if is_chapter_header(text):
            i += 1
            continue

        # Determine if this paragraph starts a new question
        fl = first_line(raw_text)
        rl = remaining_lines(raw_text)

        # Skip if first line is answer-like (continuation of previous Q)
        if is_answer_line(fl):
            i += 1
            continue

        # Skip very short fragments that are clearly not questions
        if len(fl) < 4:
            i += 1
            continue

        # This looks like a question title
        prompt = fl
        answer_parts = []

        # Add remaining lines from current paragraph
        if rl:
            answer_parts.append(rl)

        # Collect following paragraphs as answer
        j = i + 1
        while j <= end_idx:
            look_raw = paragraphs[j].text.strip('\n').replace('\r', '\n')
            look_text = clean(look_raw)
            look_fl = first_line(look_raw)

            if not look_text:
                j += 1
                continue

            # Stop at next subcategory or chapter header
            if is_subcategory_header(look_fl) or is_chapter_header(look_text):
                break

            # Stop if this looks like a new question title
            # (first line is NOT answer-like and looks like a topic title)
            if not is_answer_line(look_fl) and _is_topic_title(look_fl) and answer_parts:
                break

            # Add this paragraph's full text as answer content
            answer_parts.append(look_text)
            j += 1

        question_num += 1
        reference_answer = '\n\n'.join(answer_parts).strip()

        questions.append({
            "id": f"{current_subcat_id}-q-{question_num:03d}" if current_subcat_id else f"{type_prefix}-q-{question_num:03d}",
            "type": question_type,
            "categoryId": current_subcat_id or f"{type_prefix}-1",
            "categoryTitle": current_subcat or parent_title,
            "parentTitle": parent_title,
            "number": question_num,
            "prompt": prompt,
            "referenceAnswer": reference_answer,
        })

        i = j - 1
        i += 1

    return questions


def _is_topic_title(text: str) -> bool:
    """Check if text looks like a question/topic title (not an answer line)."""
    if not text:
        return False
    if len(text) < 6:
        return False
    if is_answer_line(text):
        return False
    if text.startswith('结论') or text.startswith('解析') or text.startswith('答案'):
        return False
    if re.match(r'^(注[：:]|注意|提示|例如|比如|补充|附[：:]|参考)', text):
        return False
    return True


def parse_short_answer(doc: Any, start_idx: int, end_idx: int) -> list[dict]:
    return parse_qa_section(doc, start_idx, end_idx, "short-answer", "简答题", "short-answer")


def parse_essay(doc: Any, start_idx: int, end_idx: int) -> list[dict]:
    """Parse essay questions, stopping after the 4th subcategory.
    The 5th+ subcategories are supplementary experimental design questions."""
    paragraphs = doc.paragraphs
    essay_end = end_idx
    subcat_count = 0
    for i in range(start_idx + 1, end_idx + 1):
        raw_text = paragraphs[i].text.strip('\n').replace('\r', '\n')
        fl = first_line(raw_text)
        if is_subcategory_header(fl):
            subcat_count += 1
            if subcat_count >= 5:
                essay_end = i - 1
                print(f"    Essay truncated at paragraph {i} after 4th subcategory (found: {clean(fl)})")
                break

    return parse_qa_section(doc, start_idx, essay_end, "essay", "分析论述题", "essay")


# ── main ─────────────────────────────────────────────────────────────

def main():
    print(f"Reading: {DOCX_PATH}")
    doc = Document(str(DOCX_PATH))

    # Phase 1: Translation tables
    print("Extracting translation terms from tables...")
    translation_questions = parse_translation_tables(doc)
    print(f"  Found {len(translation_questions)} translation terms")

    # Build paragraph map
    para_map = {}
    for idx, p in enumerate(doc.paragraphs):
        text = clean(p.text)
        if text:
            para_map[idx] = text

    # Find section boundaries
    boundaries = find_section_boundaries(para_map)

    tf_start, tf_end = boundaries.get('true-false') or (None, None)
    mc_start, mc_end = boundaries.get('multiple-choice') or (None, None)
    sa_start, sa_end = boundaries.get('short-answer') or (None, None)
    essay_start, essay_end = boundaries.get('essay') or (None, None)

    print(f"Section boundaries:")
    print(f"  True/False:      [{tf_start}, {tf_end}]")
    print(f"  Multiple Choice: [{mc_start}, {mc_end}]")
    print(f"  Short Answer:    [{sa_start}, {sa_end}]")
    print(f"  Essay:           [{essay_start}, {essay_end}]")

    all_questions = list(translation_questions)

    # Parse each section
    if tf_start is not None and tf_end is not None:
        print("\nExtracting true/false questions...")
        tf_qs = parse_true_false(doc, tf_start, tf_end)
        print(f"  Found {len(tf_qs)} questions")
        all_questions.extend(tf_qs)

    if mc_start is not None and mc_end is not None:
        print("\nExtracting multiple choice questions...")
        mc_qs = parse_multiple_choice(doc, mc_start, mc_end)
        print(f"  Found {len(mc_qs)} questions")
        all_questions.extend(mc_qs)

    if sa_start is not None and sa_end is not None:
        print("\nExtracting short answer questions...")
        sa_qs = parse_short_answer(doc, sa_start, sa_end)
        print(f"  Found {len(sa_qs)} questions")
        all_questions.extend(sa_qs)

    if essay_start is not None and essay_end is not None:
        print("\nExtracting essay questions...")
        essay_qs = parse_essay(doc, essay_start, essay_end)
        print(f"  Found {len(essay_qs)} questions")
        all_questions.extend(essay_qs)

    # Build categories
    categories_list: list[dict] = []
    seen_categories: dict[str, dict] = {}
    for q in all_questions:
        cid = q["categoryId"]
        if cid not in seen_categories:
            seen_categories[cid] = {
                "id": cid,
                "type": q["type"],
                "title": q["categoryTitle"],
                "parentTitle": q.get("parentTitle", ""),
                "questionCount": 0,
            }
            categories_list.append(seen_categories[cid])
        seen_categories[cid]["questionCount"] += 1

    output: dict[str, Any] = {
        "generatedAt": datetime.now().isoformat(),
        "totalQuestions": len(all_questions),
        "categories": categories_list,
        "questions": all_questions,
    }

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    with open(OUTPUT_PATH, 'w', encoding='utf-8') as f:
        json.dump(output, f, ensure_ascii=False, indent=2)

    print(f"\nDone: {len(all_questions)} questions written to {OUTPUT_PATH}")

    type_counts: dict[str, int] = {}
    for q in all_questions:
        t = q["type"]
        type_counts[t] = type_counts.get(t, 0) + 1
    print("Question type summary:")
    for t, c in sorted(type_counts.items()):
        print(f"  {t}: {c}")


if __name__ == '__main__':
    main()
